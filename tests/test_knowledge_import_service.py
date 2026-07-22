import subprocess
import time
from collections.abc import Callable
from pathlib import Path

import pytest
from docx import Document
from PIL import Image, ImageDraw

from xenix.config import ensure_app_dirs, get_app_paths
from xenix.services.artifact_service import ArtifactService
from xenix.services.knowledge_derivation_service import KnowledgeDerivationService
from xenix.services.knowledge_import_service import KnowledgeImportService, _find_libreoffice
from xenix.services.knowledge_pipeline import FormatNormalizer
from xenix.services.knowledge_service import KnowledgeService
from xenix.services.storage import StorageBootstrapService


@pytest.fixture
def knowledge_runtime(monkeypatch, tmp_path: Path):
    monkeypatch.setenv("XENIX_APP_HOME", str(tmp_path / "xenix-home"))
    paths = ensure_app_dirs(get_app_paths())
    storage = StorageBootstrapService().initialize(paths)
    knowledge = KnowledgeService(storage.session_factory)
    started: list[tuple[KnowledgeImportService, KnowledgeDerivationService]] = []

    def create(*, ocr_service=None) -> tuple[KnowledgeImportService, KnowledgeService]:
        derivation = KnowledgeDerivationService(
            paths=paths,
            session_factory=storage.session_factory,
        )
        importer = KnowledgeImportService(
            paths=paths,
            session_factory=storage.session_factory,
            artifact_service=ArtifactService(storage.session_factory),
            normalizer=FormatNormalizer(),
            ocr_service=ocr_service,
            canonical_ready_notifier=derivation.enqueue_generation,
        )
        started.append((importer, derivation))
        return importer, knowledge

    yield create

    for importer, derivation in reversed(started):
        importer.shutdown()
        derivation.shutdown()


def test_txt_import_publishes_docling_ir_cas_and_searchable_units(
    knowledge_runtime: Callable,
    tmp_path: Path,
) -> None:
    importer, knowledge = knowledge_runtime()
    source = tmp_path / "规则.txt"
    source.write_text("华东雨季的雨具目标库存按三周平均销量计算。", encoding="utf-8-sig")

    result = importer.import_file(source)

    assert result.reused_existing is False
    assert _wait_for_lookup(knowledge, "雨具三周销量", result.document_id)
    canonical_path = Path(result.canonical_path or "")
    assert (canonical_path / "manifest.json").is_file()
    assert (canonical_path / "canonical-envelope.json.zst").is_file()
    assert (canonical_path / "docling-document.json.zst").is_file()
    assert result.canonical_ready is True


def test_same_sha_reuses_document_within_global_library(
    knowledge_runtime: Callable,
    tmp_path: Path,
) -> None:
    importer, _knowledge = knowledge_runtime()
    first = tmp_path / "first.txt"
    second = tmp_path / "second.txt"
    first.write_text("相同的知识内容。", encoding="utf-8")
    second.write_bytes(first.read_bytes())

    initial = importer.import_file(first)
    repeated = importer.import_file(second)

    assert repeated.reused_existing is True
    assert repeated.document_id == initial.document_id
    assert repeated.source_sha256 == initial.source_sha256
    assert sorted(item.attempt_number for item in importer.list_imports()) == [1, 2]


def test_docx_uses_docling_and_becomes_searchable(
    knowledge_runtime: Callable,
    tmp_path: Path,
) -> None:
    importer, knowledge = knowledge_runtime()
    docx_path = tmp_path / "经营规则.docx"
    docx = Document()
    docx.add_heading("渠道规则", level=1)
    docx.add_paragraph("会员日活动必须保持毛利率不低于百分之十八。")
    docx.save(docx_path)
    docx_result = importer.import_file(docx_path)

    assert _wait_for_lookup(knowledge, "会员日毛利率", docx_result.document_id)


def test_born_digital_pdf_uses_docling_without_builtin_ocr(
    knowledge_runtime: Callable,
    tmp_path: Path,
) -> None:
    importer, knowledge = knowledge_runtime()
    pdf_path = tmp_path / "policy.pdf"
    _write_simple_pdf(pdf_path, "Rainy season inventory uses three week demand")

    result = importer.import_file(pdf_path)

    assert _wait_for_lookup(knowledge, "three week demand", result.document_id)


def test_scanned_pdf_routes_missing_text_page_to_local_paddle_ocr(
    knowledge_runtime: Callable,
    tmp_path: Path,
) -> None:
    class FakeLocalPaddleOcr:
        def recognize(self, image_path: Path, *, output_path: Path, timeout: int = 300) -> dict:
            assert image_path.suffix == ".png"
            return {
                "protocol": 1,
                "pages": [
                    {
                        "res": {
                            "rec_texts": ["扫描页雨具补货使用三周需求"],
                            "rec_polys": [[[10, 10], [300, 10], [300, 40], [10, 40]]],
                        }
                    }
                ],
            }

    importer, knowledge = knowledge_runtime(ocr_service=FakeLocalPaddleOcr())
    pdf_path = tmp_path / "scanned.pdf"
    image = Image.new("RGB", (600, 300), "white")
    ImageDraw.Draw(image).text((40, 100), "scanned policy", fill="black")
    image.save(pdf_path, "PDF")

    result = importer.import_file(pdf_path)

    _wait_for_lookup(knowledge, "扫描页雨具补货", result.document_id)
    match = knowledge.lookup("扫描页雨具补货")[0]
    assert match.document_id == result.document_id
    assert match.locator["page"] == 1


@pytest.mark.skipif(_find_libreoffice() is None, reason="LibreOffice is not installed")
def test_legacy_doc_normalizes_through_libreoffice_then_docling(
    knowledge_runtime: Callable,
    tmp_path: Path,
) -> None:
    importer, knowledge = knowledge_runtime()
    modern_doc = tmp_path / "legacy-rule.docx"
    docx = Document()
    docx.add_paragraph("旧版文档中的渠道库存规则")
    docx.save(modern_doc)
    executable = _find_libreoffice()
    assert executable is not None
    completed = subprocess.run(
        [str(executable), "--headless", "--convert-to", "doc:MS Word 97", "--outdir", str(tmp_path), str(modern_doc)],
        capture_output=True,
        text=True,
        timeout=120,
        check=False,
    )
    assert completed.returncode == 0

    doc_result = importer.import_file(tmp_path / "legacy-rule.doc")

    assert _wait_for_lookup(knowledge, "渠道库存规则", doc_result.document_id)


def _wait_for_lookup(
    knowledge: KnowledgeService,
    query: str,
    document_id: str,
    *,
    timeout: float = 30.0,
) -> bool:
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        if document_id in {match.document_id for match in knowledge.lookup(query)}:
            return True
        time.sleep(0.02)
    raise AssertionError("Knowledge derivation did not publish searchable Units.")


def _write_simple_pdf(path: Path, text: str) -> None:
    content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        f"<< /Length {len(content)} >>\nstream\n".encode() + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    payload = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{index} 0 obj\n".encode())
        payload.extend(body)
        payload.extend(b"\nendobj\n")
    xref = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode())
    payload.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    )
    path.write_bytes(payload)
