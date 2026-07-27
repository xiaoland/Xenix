"""Repeatable legacy-DOC normalization comparison for the Knowledge task packet."""

from __future__ import annotations

from collections import Counter
from hashlib import sha256
import json
from pathlib import Path
import shutil
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_BREAK
from PIL import Image, ImageDraw
import pikepdf

PROJECT_ROOT = Path(__file__).resolve().parents[3]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from xenix.services.knowledge_docling import convert_document


MARKERS = (
    "经营备忘 Q7",
    "雨具目标库存采用最近三周平均销量",
    "补货量必须扣除当前库存",
    "U100",
    "R200",
    "Q7-REF-2026",
)


def main() -> int:
    output_root = PROJECT_ROOT / "build" / "knowledge-doc-fidelity-spike"
    if output_root.exists():
        shutil.rmtree(output_root)
    output_root.mkdir(parents=True)

    seed_docx = output_root / "seed.docx"
    figure = output_root / "figure.png"
    _write_fixture(seed_docx, figure)

    source_dir = output_root / "source"
    source_dir.mkdir()
    _libreoffice_convert(seed_docx, source_dir, "doc:MS Word 97", output_root / "profile-seed")
    source_doc = source_dir / "seed.doc"
    if not source_doc.is_file():
        raise RuntimeError("LibreOffice did not produce the legacy DOC fixture.")

    docx_dir = output_root / "docx-route"
    pdf_dir = output_root / "pdf-route"
    docx_dir.mkdir()
    pdf_dir.mkdir()
    _libreoffice_convert(
        source_doc,
        docx_dir,
        "docx:Office Open XML Text",
        output_root / "profile-docx",
    )
    _libreoffice_convert(
        source_doc,
        pdf_dir,
        "pdf:writer_pdf_Export",
        output_root / "profile-pdf",
    )

    normalized_docx = docx_dir / "seed.docx"
    normalized_pdf = pdf_dir / "seed.pdf"
    report = {
        "schema": "xenix.knowledge-doc-fidelity-spike/v1",
        "runtime": {
            "python": sys.version.split()[0],
            "libreoffice": _libreoffice_version(),
        },
        "source": _file_identity(source_doc),
        "routes": {
            "docx": _inspect_route(normalized_docx, source_format="docx"),
            "pdf": _inspect_route(normalized_pdf, source_format="pdf"),
        },
    }
    report_path = output_root / "result.json"
    report_path.write_text(
        json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True),
        encoding="utf-8",
    )
    print(report_path)
    print(json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


def _write_fixture(path: Path, figure_path: Path) -> None:
    image = Image.new("RGB", (640, 180), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((4, 4, 635, 175), outline="navy", width=5)
    draw.text((40, 70), "FIGURE Q7 - SEASONAL STOCK", fill="navy")
    image.save(figure_path)

    document = Document()
    document.sections[0].header.paragraphs[0].text = "经营备忘 Q7"
    document.sections[0].footer.paragraphs[0].text = "内部经验 | Q7-REF-2026"
    document.add_heading("华东雨季补货经验", level=1)
    document.add_paragraph("雨具目标库存采用最近三周平均销量。")
    document.add_paragraph("补货量必须扣除当前库存，并将负值归零。")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ("SKU", "类别", "安全下限"), strict=True):
        cell.text = value
    for values in (("U100", "雨伞", "130"), ("R200", "雨衣", "75")):
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.text = value
    document.add_picture(str(figure_path))
    document.add_paragraph("图 1：季节性库存边界示意。")
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("例外与引用", level=2)
    document.add_paragraph("停售商品不得补货；来源标识 Q7-REF-2026。")
    document.save(path)


def _libreoffice_convert(source: Path, output: Path, filter_spec: str, profile: Path) -> None:
    executable = _libreoffice_executable()
    completed = subprocess.run(
        [
            str(executable),
            "--headless",
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--convert-to",
            filter_spec,
            "--outdir",
            str(output),
            str(source),
        ],
        cwd=output,
        capture_output=True,
        text=True,
        check=False,
        timeout=120,
    )
    if completed.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {completed.returncode}")


def _libreoffice_executable() -> Path:
    candidates = (
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        r"C:\Program Files\LibreOffice\program\soffice.com",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
    )
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return Path(candidate)
    raise RuntimeError("LibreOffice is unavailable.")


def _libreoffice_version() -> str:
    completed = subprocess.run(
        [str(_libreoffice_executable()), "--version"],
        capture_output=True,
        text=True,
        check=False,
        timeout=30,
    )
    return (completed.stdout or completed.stderr).strip()


def _inspect_route(path: Path, *, source_format: str) -> dict[str, object]:
    document = convert_document(path, source_format=source_format)
    exported = document.export_to_text()
    labels = Counter(
        str(getattr(getattr(item, "label", None), "value", getattr(item, "label", "")))
        for item, _level in document.iterate_items()
    )
    page_numbers = sorted(
        {
            int(page_no)
            for item, _level in document.iterate_items()
            for provenance in list(getattr(item, "prov", ()) or ())
            if (page_no := getattr(provenance, "page_no", None)) is not None
        }
    )
    tables = [table.export_to_markdown(doc=document) for table in document.tables]
    result: dict[str, object] = {
        **_file_identity(path),
        "text_characters": len(exported),
        "marker_recall": {marker: marker in exported for marker in MARKERS},
        "label_counts": dict(sorted(labels.items())),
        "table_count": len(document.tables),
        "picture_count": len(document.pictures),
        "table_markdown": tables,
        "page_numbers": page_numbers,
    }
    if source_format == "pdf":
        with pikepdf.Pdf.open(path) as pdf:
            result["physical_pages"] = len(pdf.pages)
    return result


def _file_identity(path: Path) -> dict[str, object]:
    payload = path.read_bytes()
    return {
        "file": path.name,
        "bytes": len(payload),
        "sha256": sha256(payload).hexdigest(),
    }


if __name__ == "__main__":
    raise SystemExit(main())
